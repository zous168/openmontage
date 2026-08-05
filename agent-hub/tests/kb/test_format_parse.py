"""FR-KB-14 · 上传格式解析."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from plugins.mxai.kb.format_parse import _denoise_pdf_text, _looks_like_pdf_internals, parse_upload


def test_rejects_pdf_internals_as_utf8() -> None:
    garbage = "/Type /Font /Subtype /Type3 /CharProcs << /g0 514 0 R >> /Widths [ 0 1000 ]"
    assert _looks_like_pdf_internals(garbage)


def test_denoise_spaced_cjk() -> None:
    raw = "政 企 行 业 Agent 研 究 报 告"
    clean = _denoise_pdf_text(raw)
    assert "政企行业" in clean
    assert "Agent" in clean


def test_parse_pdf_real_file_if_present() -> None:
    sample = Path(r"C:\Users\zhaoh\Desktop\政企行业Agent研究报告.pdf")
    if not sample.is_file():
        pytest.skip("sample PDF not on disk")
    parsed = parse_upload(sample.name, sample.read_bytes())
    assert parsed.parse_status == "ok"
    assert parsed.source_format == "pdf"
    assert "Agent" in parsed.text or "agent" in parsed.text.lower()
    assert not _looks_like_pdf_internals(parsed.text)
    assert "/CharProcs" not in parsed.text


def test_parse_txt_utf8() -> None:
    parsed = parse_upload("note.txt", "你好，知识库。".encode("utf-8"))
    assert parsed.parse_status == "ok"
    assert "知识库" in parsed.text


def test_parse_docx() -> None:
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("常见问题", level=1)
    doc.add_paragraph("支持本地离线部署，不上云。")
    doc.add_paragraph("可按席位授权。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "版本"
    table.rows[0].cells[1].text = "价格"
    table.rows[1].cells[0].text = "基础版"
    table.rows[1].cells[1].text = "面议"
    doc.save(buf)

    parsed = parse_upload("常见问题FAQ.docx", buf.getvalue())
    assert parsed.parse_status == "ok"
    assert parsed.source_format == "docx"
    assert "常见问题" in parsed.text
    assert "本地离线" in parsed.text
    assert "基础版" in parsed.text


def test_parse_xlsx() -> None:
    from openpyxl import Workbook

    buf = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "报价"
    ws.append(["版本", "价格"])
    ws.append(["专业版", "9980"])
    wb.save(buf)

    parsed = parse_upload("报价方案2025.xlsx", buf.getvalue())
    assert parsed.parse_status == "ok"
    assert parsed.source_format == "xlsx"
    assert "报价" in parsed.text
    assert "专业版" in parsed.text


def test_rejects_legacy_doc() -> None:
    parsed = parse_upload("old.doc", b"\x00" * 100)
    assert parsed.parse_status == "failed"
    assert "docx" in parsed.warning
